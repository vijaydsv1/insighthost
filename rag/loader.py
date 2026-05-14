import os
import json
import pandas as pd

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


# =========================================================
# Base Knowledge Path
# =========================================================
DATA_PATH = "knowledge_base"

DOCUMENTS_PATH = os.path.join(DATA_PATH, "documents")
JSON_PATH = os.path.join(DATA_PATH, "json")
CSV_PATH = os.path.join(DATA_PATH, "csv")
PDF_PATH = os.path.join(DATA_PATH, "pdfs")

IMAGES_PATH = os.path.join(DATA_PATH, "images")
VIDEOS_PATH = os.path.join(DATA_PATH, "videos")


# =========================================================
# Helper Functions
# =========================================================
def find_related_image(file_stem):

    for image_file in os.listdir(IMAGES_PATH):

        if image_file.startswith(file_stem):

            return f"/media/images/{image_file}"

    return None


def find_related_video(file_stem):

    for video_file in os.listdir(VIDEOS_PATH):

        if video_file.startswith(file_stem):

            return f"/media/videos/{video_file}"

    return None


# =========================================================
# Load Documents
# =========================================================
def load_documents():

    documents = []

    # =====================================================
    # TEXT FILES
    # =====================================================
    if os.path.exists(DOCUMENTS_PATH):

        for file_name in os.listdir(DOCUMENTS_PATH):

            file_path = os.path.join(
                DOCUMENTS_PATH,
                file_name
            )

            try:

                if file_name.lower().endswith(".txt"):

                    with open(
                        file_path,
                        "r",
                        encoding="utf-8"
                    ) as f:

                        text = f.read().strip()

                        if text:

                            file_stem = os.path.splitext(
                                file_name
                            )[0]

                            image = find_related_image(
                                file_stem
                            )

                            video = find_related_video(
                                file_stem
                            )

                            documents.append(
                                Document(
                                    page_content=text,

                                    metadata={

                                        "source": file_name,

                                        "title": file_stem,

                                        "type": "text",

                                        "image": image,

                                        "video": video
                                    }
                                )
                            )

            except Exception as e:

                print(f"Error loading TXT {file_name}: {e}")

    # =====================================================
    # JSON FILES
    # =====================================================
    if os.path.exists(JSON_PATH):

        for file_name in os.listdir(JSON_PATH):

            file_path = os.path.join(
                JSON_PATH,
                file_name
            )

            try:

                if file_name.lower().endswith(".json"):

                    with open(
                        file_path,
                        "r",
                        encoding="utf-8"
                    ) as f:

                        data = json.load(f)

                        if isinstance(data, list):

                            for item in data:

                                category = item.get(
                                    "category", ""
                                )

                                pillar = item.get(
                                    "pillar", ""
                                )

                                service = item.get(
                                    "service", ""
                                )

                                content = item.get(
                                    "content", ""
                                )

                                source = item.get(
                                    "source",
                                    file_name
                                )

                                image = item.get(
                                    "image",
                                    None
                                )

                                video = item.get(
                                    "video",
                                    None
                                )

                                text = f"""
Category: {category}
Pillar: {pillar}
Service: {service}

{content}
""".strip()

                                if text:

                                    documents.append(
                                        Document(
                                            page_content=text,

                                            metadata={
                                                "category": category,
                                                "pillar": pillar,
                                                "service": service,
                                                "source": source,
                                                "type": "json",
                                                "image": image,
                                                "video": video
                                            }
                                        )
                                    )

            except Exception as e:

                print(f"Error loading JSON {file_name}: {e}")

    # =====================================================
    # PDF FILES
    # =====================================================
    if os.path.exists(PDF_PATH):

        for file_name in os.listdir(PDF_PATH):

            file_path = os.path.join(
                PDF_PATH,
                file_name
            )

            try:

                if file_name.lower().endswith(".pdf"):

                    reader = PdfReader(file_path)

                    for page in reader.pages:

                        text = page.extract_text()

                        if text:

                            documents.append(
                                Document(
                                    page_content=text.strip(),

                                    metadata={
                                        "source": file_name,
                                        "type": "pdf"
                                    }
                                )
                            )

            except Exception as e:

                print(f"Error loading PDF {file_name}: {e}")

    # =====================================================
    # CSV FILES
    # =====================================================
    if os.path.exists(CSV_PATH):

        for file_name in os.listdir(CSV_PATH):

            file_path = os.path.join(
                CSV_PATH,
                file_name
            )

            try:

                if file_name.lower().endswith(".csv"):

                    df = pd.read_csv(
                        file_path,
                        on_bad_lines="skip"
                    )

                    for _, row in df.iterrows():

                        text = " ".join(
                            map(str, row.values)
                        ).strip()

                        if text:

                            documents.append(
                                Document(
                                    page_content=text,

                                    metadata={
                                        "source": file_name,
                                        "type": "csv"
                                    }
                                )
                            )

            except Exception as e:

                print(f"Error loading CSV {file_name}: {e}")

    # =====================================================
    # Chunking
    # =====================================================
    print(f"Loaded {len(documents)} documents")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )

    split_documents = splitter.split_documents(
        documents
    )

    print(f"Created {len(split_documents)} chunks")

    return split_documents